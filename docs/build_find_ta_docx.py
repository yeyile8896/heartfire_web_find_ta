from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("爱怎么翻译_寻找那个TA_项目流程规划.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(28, 37, 50)
MUTED = RGBColor(95, 105, 120)
ORANGE = RGBColor(143, 45, 18)
LIGHT_GRAY = "F2F4F7"
LIGHT_ORANGE = "FFF3E7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F7F9FC"
BORDER = "D0D5DD"


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=11, color=INK, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def style_para(paragraph, before=0, after=6, line=1.10, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6, line=1.10, align=None):
    p = doc.add_paragraph()
    style_para(p, before=before, after=after, line=line, align=align)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        style_para(p, before=16, after=8, line=1.10)
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        style_para(p, before=12, after=6, line=1.10)
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        style_para(p, before=8, after=4, line=1.10)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    style_para(p, before=0, after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    style_para(p, before=0, after=4, line=1.167)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=INK, size=10.5, fill=None, align=None):
    cell.text = ""
    if fill:
        shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    style_para(p, before=0, after=0, line=1.10, align=align)
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def set_table_borders(table, color=BORDER, size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=10):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_cell_margins(table)
    set_table_width(table, widths)
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=INK, size=font_size, fill=header_fill)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if index == 0 and len(headers) > 2 else None
            set_cell_text(row.cells[index], value, size=font_size, align=align)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, title, body, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, color="D6DAE1", size="4")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    style_para(p, before=0, after=2, line=1.10)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=ORANGE)
    p2 = cell.add_paragraph()
    style_para(p2, before=0, after=0, line=1.10)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_para(doc, "", after=3)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("《爱，怎么翻译》营会互动平台｜寻找那个 TA")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("项目流程规划")
    set_run_font(fr, size=9, color=MUTED)
    return doc


def add_cover(doc):
    add_para(doc, "项目流程规划文档", size=12, bold=True, color=ORANGE, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "《爱，怎么翻译》营会互动平台",
        size=25,
        bold=True,
        color=INK,
        after=6,
        line=1.05,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "寻找那个 TA",
        size=22,
        bold=True,
        color=BLUE,
        after=16,
        line=1.05,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(
        doc,
        "匿名配对、线下相认、双人任务与积分排行榜网页平台",
        size=12.5,
        color=MUTED,
        after=28,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_table(
        doc,
        ["项目项", "内容"],
        [
            ("项目名称", "《爱，怎么翻译》营会互动平台"),
            ("核心活动", "寻找那个 TA"),
            ("目标人数", "约 60 名营员"),
            ("使用方式", "手机浏览器扫码进入，无需下载 App"),
            ("当前版本", "v0.1 流程规划与 MVP 功能说明"),
            ("文档日期", "2026-06-18"),
            ("文档用途", "DDL 提交、项目讲解、后续功能细化依据"),
        ],
        [1.65, 4.85],
        header_fill=LIGHT_BLUE,
        font_size=10.5,
    )

    add_callout(
        doc,
        "一句话概括",
        "这是一个面向青年营会的互动网页平台：营员扫码进入房间，匿名配对寻找神秘伙伴，完成会合确认与双人挑战，系统实时统计积分并在主持人大屏展示现场进度。",
        fill=LIGHT_ORANGE,
    )

    doc.add_page_break()


def add_content(doc):
    add_heading(doc, "1. 项目背景与目标", 1)
    add_para(
        doc,
        "「寻找那个 TA」是营会前期的破冰活动。它希望把营员从“先知道名字，再认识人”的习惯中带出来，改成先通过兴趣、穿搭、暗号和互动任务认识彼此，再发现对方是谁。",
    )
    add_para(
        doc,
        "本项目的目标是做成一个手机浏览器可访问的活动网页平台，让约 60 名营员可以在活动现场完成扫码加入、匿名配对、寻找搭档、双人挑战和积分排名。",
    )
    add_bullet(doc, "降低陌生感：用匿名线索代替直接介绍，让第一轮交流更自然。")
    add_bullet(doc, "增强现场互动：通过大屏、匿名代号墙和排行榜制造参与感。")
    add_bullet(doc, "减少人工统计：配对、确认、积分和任务状态尽量由系统自动记录。")
    add_bullet(doc, "便于主持人掌控节奏：后台可查看真实配对表和每组进度。")

    add_heading(doc, "2. 项目范围", 1)
    add_table(
        doc,
        ["范围", "说明"],
        [
            ("已包含", "扫码加入、房间码、随机配对、营员个人页、双方确认会合、共同点挑战、积分排行榜、主持人后台、主持人大屏。"),
            ("后续扩展", "匿名聊天、默契测试、营地探索、照片上传、工作人员审核、Excel 导出、本地数据库持久化。"),
            ("不做内容", "不做原生 App，不要求营员下载安装；现阶段不公开真实姓名、小组、手机号或社交账号。"),
        ],
        [1.45, 5.05],
        font_size=10,
    )

    doc.add_page_break()
    add_heading(doc, "3. 用户角色与权限", 1)
    add_table(
        doc,
        ["角色", "可以看到", "不能看到 / 不应公开"],
        [
            ("营员", "自己的匿名代号、神秘伙伴匿名线索、组合积分、组合任务状态。", "对方真实姓名、对方小组、后台配对表、其他组合真实信息。"),
            ("主持人", "真实姓名、小组、匿名代号、配对结果、会合状态、任务状态、积分。", "主持人链接不应公开给营员。"),
            ("主持人大屏", "房间码、二维码、匿名代号墙、人数统计、排行榜、现场动态。", "真实姓名、小组、后台配对关系、隐私信息。"),
        ],
        [1.1, 2.7, 2.7],
        header_fill=LIGHT_BLUE,
        font_size=9.5,
    )

    add_heading(doc, "4. 活动总流程", 1)
    add_table(
        doc,
        ["阶段", "名称", "主持人动作", "营员动作", "系统结果"],
        [
            ("0", "准备房间", "创建房间，打开大屏。", "等待扫码。", "生成房间码、二维码、后台链接。"),
            ("1", "扫码加入", "观察加入人数。", "填写真实姓名、小组和匿名线索。", "大屏显示匿名代号。"),
            ("2", "随机配对", "点击开始随机配对。", "等待页面刷新。", "系统生成 2 人匿名组合。"),
            ("3", "寻找 TA", "维持现场节奏。", "根据线索寻找神秘伙伴。", "营员看到对方线索但看不到真实姓名。"),
            ("4", "双方确认", "查看确认进度。", "两人都点击我已找到 TA。", "双方确认后加 10 分。"),
            ("5", "共同点挑战", "查看任务进度。", "填写 1 个共同点、1 个不同点，对方确认。", "双方确认后加 20 分。"),
            ("6", "默契测试", "开放测试。", "双方分别答题。", "系统计算默契度并加 20 分。"),
            ("7", "营地探索", "发布/审核任务。", "完成现场探索任务。", "通过后加 30 分。"),
            ("8", "排行榜收尾", "公布结果。", "查看排名。", "生成最终榜单和奖项。"),
        ],
        [0.5, 1.1, 1.5, 1.65, 1.75],
        header_fill=LIGHT_BLUE,
        font_size=8.7,
    )

    add_heading(doc, "5. 详细流程设计", 1)
    add_heading(doc, "5.1 阶段 0：主持人准备", 2)
    add_para(doc, "主持人进入后台创建房间。系统会生成房间码、营员加入二维码、主持人大屏链接和主持人后台链接。活动开始前，主持人应将大屏投影到现场，让营员扫码加入。")
    add_bullet(doc, "后台链接只给主持人使用，不能公开投屏。")
    add_bullet(doc, "大屏链接可以公开投屏，但只显示匿名信息。")
    add_bullet(doc, "二维码最好同时显示房间码，方便手机扫码失败时手动输入。")

    add_heading(doc, "5.2 阶段 1：营员扫码加入", 2)
    add_para(doc, "营员用手机浏览器扫描二维码进入加入页面。加入页面收集后台识别所需信息和匿名寻找所需线索。")
    add_table(
        doc,
        ["字段", "用途", "可见范围"],
        [
            ("真实姓名", "后台识别营员身份。", "主持人后台可见，营员和大屏不可见。"),
            ("小组", "方便主持人管理与后续分组统计。", "主持人后台可见，营员和大屏不可见。"),
            ("匿名代号", "大屏和营员页面显示的公开身份。", "营员、大屏、后台可见。"),
            ("兴趣爱好", "帮助神秘伙伴寻找自己。", "神秘伙伴可见。"),
            ("今日穿搭", "帮助线下观察寻找。", "神秘伙伴可见。"),
            ("确认暗号", "双方现实中确认身份。", "神秘伙伴可见。"),
            ("给 TA 的一句话", "增加温度，可选。", "神秘伙伴可见。"),
        ],
        [1.35, 3.0, 2.15],
        font_size=9.5,
    )

    add_heading(doc, "5.3 阶段 2：随机配对", 2)
    add_para(doc, "主持人点击「开始随机配对」后，系统将当前房间内的营员随机分成组合。默认以 2 人一组为主；如果人数为单数，后续需要确认是生成 3 人组，还是让最后一人等待下一轮。")
    add_callout(
        doc,
        "建议规则",
        "为了现场体验稳定，建议单数时允许最后 3 人组成三人组；否则最后一位营员会没有搭档，活动体验容易被打断。",
        fill=CALLOUT,
    )

    add_heading(doc, "5.4 阶段 3：匿名寻找", 2)
    add_para(doc, "配对完成后，营员个人页面会显示神秘伙伴的匿名代号和线索。营员需要在现场通过观察、交流和暗号确认找到对方。")
    add_table(
        doc,
        ["允许分享", "禁止分享"],
        [
            ("兴趣爱好、穿搭细节、指定暗号、模糊经历。", "真实姓名、照片、房间号、手机号、社交媒体账号、明显身份信息。"),
        ],
        [3.25, 3.25],
        header_fill=LIGHT_ORANGE,
        font_size=10,
    )

    add_heading(doc, "5.5 阶段 4：双方确认会合", 2)
    add_para(doc, "两位营员现实中找到彼此后，需要互相确认暗号并完成击掌。随后双方都在自己的页面点击「我已找到 TA」。")
    add_table(
        doc,
        ["确认状态", "页面显示", "是否加分"],
        [
            ("0/2", "等待双方确认。", "不加分。"),
            ("1/2", "已确认，等 TA；另一方看到 TA 已确认。", "不加分。"),
            ("2/2", "已成功会合。", "加 10 分。"),
        ],
        [1.0, 4.0, 1.5],
        font_size=10,
    )

    doc.add_page_break()
    add_heading(doc, "5.6 阶段 5：共同点挑战", 2)
    add_para(doc, "共同点挑战在双方确认会合后解锁。它的目标是让刚找到彼此的两位营员进行一次更深入但轻量的面对面交流。")
    add_bullet(doc, "解锁条件：双方都已确认会合。")
    add_bullet(doc, "必填内容：1 个共同点、1 个不同点。")
    add_bullet(doc, "可选内容：最意外的发现。")
    add_bullet(doc, "完成方式：一方填写提交，另一方确认同一份答案。")
    add_bullet(doc, "得分：双方确认后获得 20 分。")
    add_number(doc, "任意一方填写共同点挑战表单。")
    add_number(doc, "提交后，提交者自动确认，页面显示等待 TA 确认。")
    add_number(doc, "另一方页面显示同一份答案。")
    add_number(doc, "另一方点击确认后，任务完成。")
    add_number(doc, "积分从 10 / 80 变为 30 / 80，大屏排行榜同步更新。")

    add_heading(doc, "5.7 阶段 6：默契测试", 2)
    add_para(doc, "默契测试为后续计划功能。建议由系统随机生成 10 道二选一题，双方分别在手机上作答，作答结束后系统计算默契度百分比。")
    add_table(
        doc,
        ["题目示例", "选项 A", "选项 B"],
        [
            ("饮品偏好", "咖啡", "奶茶"),
            ("旅行偏好", "海边", "山上"),
            ("作息偏好", "早起", "晚睡"),
            ("手机系统", "iPhone", "Android"),
            ("社交状态", "安静", "热闹"),
        ],
        [2.0, 2.25, 2.25],
        font_size=10,
    )
    add_bullet(doc, "双方都提交答案后自动完成。")
    add_bullet(doc, "系统显示默契度百分比。")
    add_bullet(doc, "完成后获得 20 分。")
    add_bullet(doc, "默契度最高组合可进入特别奖项。")

    add_heading(doc, "5.8 阶段 7：营地探索", 2)
    add_para(doc, "营地探索为后续计划功能，目的是把组合互动从手机页面带回真实营地场景。")
    add_table(
        doc,
        ["任务类型", "示例", "完成认证方式"],
        [
            ("工作人员互动", "找工作人员合照。", "上传照片或工作人员确认。"),
            ("组合互动", "找另一组搭档合照。", "上传照片。"),
            ("地点探索", "找到指定地点拍照。", "上传照片或现场扫码。"),
            ("营会标志物", "找到指定标志物。", "上传照片。"),
        ],
        [1.45, 2.55, 2.5],
        font_size=9.5,
    )
    add_bullet(doc, "完成后获得 30 分。")
    add_bullet(doc, "建议主持人后台保留审核入口，避免随便上传照片就得分。")

    add_heading(doc, "6. 积分系统", 1)
    add_table(
        doc,
        ["任务", "分数", "完成条件", "当前状态"],
        [
            ("双方确认会合", "10", "两人都点击我已找到 TA。", "已实现"),
            ("共同点挑战", "20", "一方提交，另一方确认。", "已实现"),
            ("默契测试", "20", "双方完成答题并生成默契度。", "待实现"),
            ("营地探索", "30", "完成探索任务并通过认证。", "待实现"),
            ("总分", "80", "四个任务全部完成。", "规划中"),
        ],
        [1.7, 0.75, 2.85, 1.2],
        header_fill=LIGHT_BLUE,
        font_size=9.5,
    )

    doc.add_page_break()
    add_heading(doc, "7. 页面与功能模块", 1)
    add_table(
        doc,
        ["页面", "路径示例", "核心功能"],
        [
            ("活动说明页", "/find-ta", "介绍活动规则，引导主持人或营员进入对应流程。"),
            ("营员加入页", "/find-ta/join/房间码", "扫码后填写真实姓名、小组和匿名线索。"),
            ("营员个人页", "/find-ta/room/房间码/me/营员ID", "查看神秘伙伴线索、确认会合、完成任务、查看积分。"),
            ("主持人后台", "/find-ta/host?room=房间码&token=token", "创建房间、查看真实配对表、开始配对、查看任务进度。"),
            ("主持人大屏", "/find-ta/screen/房间码", "展示二维码、匿名代号墙、排行榜和现场动态。"),
        ],
        [1.45, 2.55, 2.5],
        font_size=9,
    )

    add_heading(doc, "8. 状态设计", 1)
    add_heading(doc, "8.1 会合状态", 2)
    add_table(
        doc,
        ["状态", "说明"],
        [
            ("等待中", "还没有开始配对。"),
            ("已配对", "已生成神秘伙伴，可以开始寻找。"),
            ("等 TA 确认", "自己已确认会合，对方还没有确认。"),
            ("TA 已确认", "对方已确认会合，自己还没有确认。"),
            ("已会合", "双方都确认成功，获得 10 分。"),
        ],
        [1.6, 4.9],
        font_size=10,
    )

    doc.add_page_break()
    add_heading(doc, "8.2 共同点挑战状态", 2)
    add_table(
        doc,
        ["状态", "说明"],
        [
            ("会合后解锁", "双方还没有确认会合。"),
            ("可填写", "已会合，可以提交共同点挑战。"),
            ("等待 TA 确认", "自己已提交或确认，等待对方确认。"),
            ("请确认", "对方已提交，自己需要确认。"),
            ("已完成", "双方确认完成，获得 20 分。"),
        ],
        [1.6, 4.9],
        font_size=10,
    )

    doc.add_page_break()
    add_heading(doc, "9. 数据可见范围与隐私原则", 1)
    add_table(
        doc,
        ["数据", "营员", "主持人后台", "主持人大屏"],
        [
            ("真实姓名", "不可见", "可见", "不可见"),
            ("小组", "不可见", "可见", "不可见"),
            ("匿名代号", "可见", "可见", "可见"),
            ("兴趣/穿搭/暗号", "仅神秘伙伴可见", "可见", "不可见"),
            ("积分和任务状态", "自己组合可见", "可见", "排行榜可见"),
            ("配对关系", "只知道自己的 TA", "完整可见", "不显示真实配对"),
        ],
        [1.45, 1.45, 1.7, 1.9],
        font_size=9,
    )
    add_callout(
        doc,
        "隐私原则",
        "公开页面只显示匿名信息；真实姓名和小组只用于主持人管理，不进入大屏。后续如果加入聊天或照片上传，需要进一步加入禁词、审核和数据清理规则。",
        fill=CALLOUT,
    )

    add_heading(doc, "10. 当前实现情况", 1)
    add_table(
        doc,
        ["功能", "状态", "备注"],
        [
            ("创建房间与二维码", "已实现", "主持人可创建房间并展示二维码。"),
            ("扫码加入", "已实现", "营员手机浏览器可加入。"),
            ("匿名代号墙", "已实现", "大屏实时显示匿名代号。"),
            ("随机配对", "已实现", "支持 2 人配对，单数时当前可形成 3 人组。"),
            ("双方确认会合", "已实现", "双方确认后才加 10 分。"),
            ("共同点挑战", "已实现", "必填 1 个共同点和 1 个不同点；可选意外发现。"),
            ("排行榜", "已实现", "会合和共同点挑战完成后实时更新分数。"),
            ("主持人后台", "已实现", "可查看真实姓名、小组、配对和任务状态。"),
            ("主持人大屏", "已实现", "二维码、房间码、匿名代号墙、排行榜和动态。"),
        ],
        [2.0, 1.1, 3.4],
        header_fill=LIGHT_BLUE,
        font_size=9.2,
    )

    add_heading(doc, "11. 待开发功能", 1)
    add_table(
        doc,
        ["功能", "优先级", "说明"],
        [
            ("本地数据库持久化", "高", "当前本地测试数据主要在服务内存中，重启后可能丢失。正式活动建议改为 SQLite。"),
            ("默契测试", "高", "完成 10 道二选一题，自动计算默契度并加分。"),
            ("营地探索", "中", "生成探索任务，上传照片或工作人员审核。"),
            ("数据导出", "中", "导出用户名单、配对结果、任务完成情况和排行榜。"),
            ("匿名聊天", "可选", "如果现场允许手机聊天，可加入文字聊天和禁词规则。"),
            ("禁词检测", "可选", "防止营员直接发真实姓名、手机号或社交账号。"),
        ],
        [1.75, 1.0, 3.75],
        font_size=9.5,
    )

    add_heading(doc, "12. 技术与部署方案", 1)
    add_para(doc, "当前项目是网页方案，不是原生 App。营员只需要通过手机浏览器访问链接，适合营会现场快速使用。")
    add_table(
        doc,
        ["方案", "优点", "风险 / 注意事项"],
        [
            ("本地运行 + 同一 Wi-Fi", "不依赖云数据库，成本低，适合测试。", "手机和电脑必须在同一网络；电脑不能关机；重启可能丢数据。"),
            ("本地数据库 SQLite", "数据保存在电脑本地，免费，不担心云额度。", "需要做好备份；公网访问需要隧道或服务器。"),
            ("公网域名 + 本地服务隧道", "手机可用域名访问，仍可用本地数据库。", "依赖网络稳定性；需要配置 HTTPS/隧道。"),
            ("Firebase / Supabase", "实时同步方便，部署成熟。", "虽然免费额度通常够 60 人活动，但用户担心额度和云依赖。"),
        ],
        [1.65, 2.35, 2.5],
        font_size=8.9,
    )
    doc.add_page_break()
    add_callout(
        doc,
        "建议",
        "如果活动当天最担心费用和云额度，推荐下一步改成本地 SQLite 数据库，再配合公共域名或隧道给手机访问。这样数据主要在本机，成本可控，也便于活动后导出。",
        fill=LIGHT_ORANGE,
    )

    add_heading(doc, "13. DDL 版本交付重点", 1)
    add_para(doc, "如果今天需要提交项目细节，建议重点展示以下内容，能说明项目已经有清晰目标、完整流程和可落地实现路径。")
    add_bullet(doc, "项目不是单纯网页展示，而是完整活动互动平台。")
    add_bullet(doc, "营员、主持人和大屏三种视角已经规划清楚。")
    add_bullet(doc, "配对、会合、共同点挑战和积分已有明确规则。")
    add_bullet(doc, "隐私边界清楚：公开只显示匿名信息，真实姓名和小组只在后台可见。")
    add_bullet(doc, "后续功能路线明确：默契测试、营地探索、数据库持久化、数据导出。")

    add_heading(doc, "14. 后续需要确认的问题", 1)
    add_table(
        doc,
        ["问题", "需要决定的内容"],
        [
            ("配对规则", "是否尽量跨小组配对；是否允许三人组；迟到营员如何处理。"),
            ("匿名聊天", "是否需要网页内聊天，还是只靠线索和现场寻找。"),
            ("共同点挑战", "是否允许提交后修改；是否需要主持人审核；是否精选展示到大屏。"),
            ("默契测试", "题目数量、题库内容、是否进入特别榜单。"),
            ("营地探索", "是否上传照片、是否需要工作人员审核、任务是否随机。"),
            ("数据保存", "是否必须断网可用；是否使用本地数据库；是否导出 Excel。"),
        ],
        [1.55, 4.95],
        font_size=9.5,
    )

    doc.add_page_break()
    add_heading(doc, "15. 建议开发顺序", 1)
    add_table(
        doc,
        ["顺序", "建议动作"],
        [
            ("1", "确认最终活动规则和主持人现场流程。"),
            ("2", "确认配对规则，尤其是同组/跨组、单数、迟到营员。"),
            ("3", "将当前内存数据升级为本地 SQLite 数据库。"),
            ("4", "完成默契测试功能，并接入积分榜。"),
            ("5", "完成营地探索功能和审核方式。"),
            ("6", "加入数据导出，便于活动后保存结果。"),
            ("7", "根据公共域名确认最终部署方案。"),
            ("8", "活动前做 60 人压力测试和手机扫码测试。"),
        ],
        [0.8, 5.7],
        font_size=10,
    )


def main():
    doc = setup_document()
    add_cover(doc)
    add_content(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
